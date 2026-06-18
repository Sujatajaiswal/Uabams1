from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "UABAMS_Cloud_Data_Interface_Handoff_v2.docx"
FLOW = ROOT / "_cloud_data_flow.png"

NAVY = "123A63"
BLUE = "2E74B5"
PALE = "E8EEF5"
LIGHT = "F4F6F9"
GREEN = "E8F3EC"
AMBER = "FFF4D6"
RED = "FCE8E6"
GRAY = "5B6573"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.width = Inches(sum(widths))
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(int(widths[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for idx, value in enumerate(headers):
        shade(hdr.cells[idx], PALE)
        p = hdr.cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(value))
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(font_size)
        run.font.color.rgb = rgb(NAVY)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.name = "Calibri"
            run.font.size = Pt(font_size)
            if idx > 0 and len(str(value)) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Code Block"]
    for line in text.strip().splitlines():
        p.add_run(line + "\n")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_callout(doc, title, text, fill=LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    shade(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run(title + ": ")
    r.bold = True
    r.font.color.rgb = rgb(NAVY)
    p.add_run(text)
    set_table_widths(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def draw_flowchart():
    img = Image.new("RGB", (1600, 920), "white")
    d = ImageDraw.Draw(img)
    try:
        title = ImageFont.truetype("arialbd.ttf", 36)
        bold = ImageFont.truetype("arialbd.ttf", 24)
        body = ImageFont.truetype("arial.ttf", 21)
    except OSError:
        title = bold = body = ImageFont.load_default()

    d.text((50, 30), "UABAMS Cloud Data Flow", fill="#123A63", font=title)

    def box(x, y, w, h, heading, lines, fill):
        d.rounded_rectangle((x, y, x + w, y + h), radius=12, fill=fill, outline="#58728D", width=3)
        d.text((x + 18, y + 14), heading, fill="#123A63", font=bold)
        yy = y + 54
        for line in lines:
            for part in wrap(line, 35):
                d.text((x + 18, yy), part, fill="#243548", font=body)
                yy += 27

    def arrow(x1, y1, x2, y2, label=""):
        d.line((x1, y1, x2, y2), fill="#2E74B5", width=5)
        d.polygon([(x2, y2), (x2 - 15, y2 - 9), (x2 - 15, y2 + 9)], fill="#2E74B5")
        if label:
            d.text(((x1 + x2) // 2 - 80, y1 - 33), label, fill="#5B6573", font=body)

    box(40, 150, 250, 200, "Sensors", ["ADXL Left/Right", "Bogie accelerometer", "Encoder + GPS"], "#E8F3EC")
    box(350, 135, 300, 230, "Edge Gateway", ["Builds spatial RMS", "Builds 50 m peaks", "Stores closed session ZIP"], "#E8EEF5")
    box(735, 135, 310, 230, "Cloud Ingestion API", ["HTTPS PUT /api/v1/archive", "Validate, checksum, parse", "Idempotency by gateway/session"], "#FFF4D6")
    box(1130, 90, 400, 170, "Durable Object Storage", ["Original ZIP", "Extracted binary files", "Required for production"], "#F4F6F9")
    box(1130, 315, 400, 190, "PostgreSQL", ["Metadata and parsed records", "Alerts, thresholds, audit logs", "Dashboard query source"], "#E8F3EC")
    box(740, 570, 300, 200, "Processing", ["Threshold rules", "GPS/KM association", "Notification outbox", "TMS handoff package"], "#FCE8E6")
    box(1130, 585, 400, 170, "Consumers", ["Dashboard and API", "Railway officials (SMS)", "CRIS/TMS server"], "#E8EEF5")

    arrow(290, 250, 350, 250, "RS-485/UART")
    arrow(650, 250, 735, 250, "HTTPS ZIP")
    arrow(1045, 200, 1130, 175, "archive")
    arrow(1045, 300, 1130, 390, "records")
    d.line((930, 365, 890, 570), fill="#2E74B5", width=5)
    d.polygon([(890, 570), (882, 550), (900, 554)], fill="#2E74B5")
    arrow(1040, 670, 1130, 670, "outputs")
    d.line((840, 570, 600, 430), fill="#7A5A00", width=4)
    d.polygon([(600, 430), (621, 432), (612, 449)], fill="#7A5A00")
    d.text((430, 390), "Cloud config response (planned gateway polling)", fill="#7A5A00", font=body)

    img.save(FLOW)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb("243548")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in (
        ("Title", 24, NAVY, 0, 8),
        ("Subtitle", 12, GRAY, 0, 16),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, NAVY, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(8)
    code.font.color.rgb = rgb("1F2933")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0
    p_pr = code.element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)

    header = section.header.paragraphs[0]
    header.text = "UABAMS | Cloud Data Interface Handoff"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = rgb(GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Reference implementation | 18 June 2026   |   Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)
    for r in footer.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = rgb(GRAY)


def build():
    draw_flowchart()
    doc = Document()
    configure_doc(doc)

    p = doc.add_paragraph(style="Title")
    p.add_run("UABAMS Cloud Data Interface and Implementation Handoff")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Gateway-to-cloud data contract, storage model, processing workflow, and CRIS/TMS outputs")
    add_table(doc, ["Document", "Value"], [
        ("Audience", "Railman gateway team, cloud implementers, CRIS/TMS integrators, reviewers"),
        ("Reference implementation", "FastAPI + PostgreSQL + React, deployed on Render"),
        ("Primary gateway ICD", "UABAMS Cloud Interface Document v1.2, Frozen, 09 June 2026"),
        ("Document status", "Implementation handoff for review; production decisions clearly identified"),
        ("Verified on", "18 June 2026"),
    ], [1.75, 4.75], 9)
    add_callout(doc, "Executive position", "The cloud reference flow is operational for PostgreSQL, dashboard APIs, thresholds, alerts, and export generation. Production rollout still requires durable object storage, gateway authentication, a proven real ICD ZIP acceptance test, and final CRIS/SMS connection details.", AMBER)
    add_callout(doc, "Terminology correction", "MMD and MDB are different. MMD means Maximum Moving Dimension and is a hardware mounting/clearance envelope. MDB means Microsoft Access database file and is the preferred TMS data container mentioned by the specification.", RED)

    doc.add_heading("1. Purpose and scope", level=1)
    doc.add_paragraph("This document tells an independent implementation team exactly what the UABAMS cloud receives, how the payload is validated and stored, what the cloud processes, what it sends onward, and how a gateway connects. It separates the frozen gateway ICD from demonstration-only interfaces and from integrations that depend on external railway infrastructure.")
    add_bullet(doc, "In scope: archive upload, JSON metadata, binary data types, PostgreSQL structure, processing rules, dashboard/API access, cloud-to-gateway configuration, notifications, and CRIS/TMS export.")
    add_bullet(doc, "Out of scope for cloud software: sensor mounting, MMD/SOD hardware compliance, coach wiring certification, cellular/APN procurement, and final CRIS endpoint ownership.")

    doc.add_heading("2. Current verified status", level=1)
    add_table(doc, ["Check", "Result", "Evidence / limitation"], [
        ("Backend health", "PASS", "GET /health returned HTTP 200"),
        ("Cloud database", "PASS", "GET /health/db returned PostgreSQL with live rows"),
        ("Dashboard API", "PASS", "GET /api/v1/dashboard returned HTTP 200"),
        ("Gateway and alert APIs", "PASS", "GET endpoints returned HTTP 200"),
        ("Real ICD ZIP on live cloud", "NOT PROVEN", "Parser exists, but no production gateway ZIP acceptance evidence is available"),
        ("Permanent archive storage", "GAP", "Current ARCHIVE_STORAGE_DIR is server filesystem; use object storage or persistent disk"),
        ("Realtime gateway alert POST", "GAP", "ICD says gateway HTTP sender is a stub; cloud currently generates alerts during archive/JSON ingestion"),
        ("Direct CRIS/TMS transfer", "READY, NOT CONNECTED", "HTTP delivery code exists; CRIS URL/token must be supplied"),
    ], [1.65, 1.15, 3.7], 8.5)
    add_callout(doc, "Important", "Do not describe the whole system as 'perfectly complete'. The deployed reference API is healthy, but the end-to-end railway production chain cannot be accepted until a real gateway archive, durable storage, security controls, notification provider, and CRIS receiving system are tested together.", RED)

    doc.add_heading("3. End-to-end cloud data flow", level=1)
    doc.add_picture(str(FLOW), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 1. Recommended production flow. The blue path is implemented in the reference application. The gold return path requires gateway polling support.", style="Caption")
    add_table(doc, ["Step", "Owner", "Action", "Result"], [
        ("1", "Sensor nodes", "Capture acceleration, encoder count, speed, and GPS", "Raw frames"),
        ("2", "Gateway", "Validate CRC; compute approximately 25 cm RMS and 50 m peaks", "Session files"),
        ("3", "Gateway", "Close session and create one ZIP", "Authoritative archive"),
        ("4", "Gateway -> Cloud", "HTTPS PUT raw ZIP to /api/v1/archive", "Upload request"),
        ("5", "Cloud", "Persist, checksum, validate identity/version/files", "Accepted, duplicate, incomplete, or rejected"),
        ("6", "Cloud", "Parse fixed binary records and store PostgreSQL rows", "Queryable structured data"),
        ("7", "Cloud", "Evaluate thresholds, add GPS/KM context, queue notifications", "Alerts and audit trail"),
        ("8", "Cloud", "Serve dashboard/API and generate TMS package", "Operator and CRIS outputs"),
    ], [0.45, 1.15, 3.15, 1.75], 8.2)

    doc.add_heading("4. Gateway connection to cloud", level=1)
    add_table(doc, ["Parameter", "Required value"], [
        ("Network", "4G/5G modem, railway private APN/VPN, or approved Internet path"),
        ("Application protocol", "HTTP/1.1; production URL must use HTTPS/TLS"),
        ("Method and path", "PUT https://<cloud-host>/api/v1/archive"),
        ("Request body", "Raw ZIP bytes; not multipart and not individual files"),
        ("Required headers", "Content-Length; Content-Type: application/zip; X-Archive-Name recommended"),
        ("Success", "HTTP 200 or 201 only after durable archive persistence"),
        ("Validation failure", "HTTP 422 with JSON diagnostic; gateway retains and retries"),
        ("Retry", "Up to 5 attempts, 60 s backoff; periodic rescan about every 3600 s"),
        ("Identity", "gatewayId + sessionName is the idempotency identity"),
        ("Authentication", "Production requirement: per-gateway bearer token or mTLS; not enforced by reference code"),
    ], [1.8, 4.7], 8.7)
    add_code(doc, """PUT /api/v1/archive HTTP/1.1
Host: cloud.example.railway
Content-Type: application/zip
Content-Length: <bytes>
X-Archive-Name: GW_BOGIE_001__TRAIN_07__SESSION_20260609_083015.zip
Authorization: Bearer <gateway-specific-token>

<raw ZIP bytes>""")
    doc.add_heading("4.1 Response contract", level=2)
    add_code(doc, """HTTP/1.1 200 OK
Content-Type: application/json

{
  "status": "success",
  "archiveId": 2451,
  "alertsGenerated": 2,
  "message": "Session stored"
}""")
    add_table(doc, ["HTTP", "Meaning", "Gateway action"], [
        ("200/201", "Durably accepted, or duplicate already stored", "Delete local ZIP/session"),
        ("400", "Malformed/empty ZIP or JSON", "Retain and retry; inspect fault"),
        ("401/403", "Authentication/authorization failed", "Retain; correct credential"),
        ("422", "Filename, metadata, schema, or field validation failed", "Retain; quarantine/operator review"),
        ("500/503", "Cloud/storage/database failure", "Retain and retry"),
    ], [0.75, 2.85, 2.9], 8.5)

    doc.add_heading("5. What the cloud receives", level=1)
    doc.add_heading("5.1 Production gateway ZIP", level=2)
    add_table(doc, ["Path inside ZIP", "Format", "Purpose", "Record rule"], [
        ("session_metadata.json", "UTF-8 JSON", "Identity and session state", "One object"),
        ("rms/rms_25cm.bin", "Little-endian binary", "Spatial RMS acceleration", "66 bytes/record"),
        ("peak/peak_50m.bin", "Little-endian binary", "Processed peak windows", "302 bytes/record"),
        ("faults/faults.bin", "Little-endian binary", "Gateway and node faults", "75 bytes/record"),
        ("raw/adxl_left.bin", "Length-prefixed binary", "Forensic left sensor frames", "uint32 length + N bytes"),
        ("raw/adxl_right.bin", "Length-prefixed binary", "Forensic right sensor frames", "uint32 length + N bytes"),
        ("raw/bogie.bin", "Length-prefixed binary", "Forensic bogie frames", "uint32 length + N bytes"),
        ("raw/encoder.bin", "Length-prefixed binary", "Forensic encoder frames", "uint32 length + N bytes"),
    ], [1.55, 1.35, 2.3, 1.3], 8)

    doc.add_heading("5.2 session_metadata.json data types", level=2)
    add_table(doc, ["Field", "JSON type", "Required", "Rule / example"], [
        ("schemaVersion", "string", "Yes", "Must be '1.0'"),
        ("gatewayId", "string", "Yes", "Logical ID; e.g. GW_BOGIE_001"),
        ("gatewaySerial", "string", "Yes", "Physical unit serial"),
        ("trainId", "string", "Yes", "e.g. TRAIN_07"),
        ("firmwareVersion", "string", "Yes", "Semantic/version text"),
        ("gatewaySoftware", "string", "Yes", "Literal 'UABAMS Gateway'"),
        ("sessionName", "string", "Yes", "SESSION_YYYYMMDD_HHMMSS"),
        ("sessionStatus", "string enum", "Yes", "active | closed; process complete only when closed"),
        ("createdUtc", "ISO-8601 string", "Yes", "UTC: YYYY-MM-DDTHH:MM:SSZ"),
    ], [1.45, 1.15, 0.75, 3.15], 8.2)
    add_code(doc, """{
  "schemaVersion": "1.0",
  "gatewayId": "GW_BOGIE_001",
  "gatewaySerial": "UABAMS-2026-0001",
  "trainId": "TRAIN_07",
  "firmwareVersion": "1.0.0",
  "gatewaySoftware": "UABAMS Gateway",
  "sessionName": "SESSION_20260609_083015",
  "sessionStatus": "closed",
  "createdUtc": "2026-06-09T08:30:15Z"
}""")

    doc.add_heading("5.3 Demonstration JSON (not the Railman production contract)", level=2)
    doc.add_paragraph("The reference dashboard can create synthetic data through POST /api/v1/archive. This is useful for UI and API testing but does not replace the ZIP contract.")
    add_code(doc, """{
  "gatewayId": "GW001",
  "trainId": "TRAIN07",
  "sessionId": "DEMO-20260618-001",
  "timestamp": "2026-06-18T09:41:14Z",
  "route": "Bangalore-Chennai",
  "gps": {"lat": 13.0827, "lon": 80.2707},
  "speedKmph": 95.2,
  "axleData": [
    {"axleId": "AX01", "verticalG": 27.5, "lateralG": 16.5,
     "rms": 15.13, "peak": 27.5}
  ]
}""")

    doc.add_heading("6. Binary data formats", level=1)
    doc.add_heading("6.1 RMS spatial record: 66 bytes", level=2)
    add_table(doc, ["Offset", "Field(s)", "Wire type", "Units / rule"], [
        ("0", "masterCount", "uint64 LE", "Encoder count"),
        ("8", "positionMm", "int32 LE", "Millimetres"),
        ("12", "latitude", "IEEE754 float64 LE", "WGS-84 degrees"),
        ("20", "longitude", "IEEE754 float64 LE", "WGS-84 degrees"),
        ("28", "gpsValid", "uint8/bool", "1 valid; 0 estimated/lost"),
        ("29", "validMask", "uint8 bitmask", "0x01 AL, 0x02 AR, 0x04 BG"),
        ("30..65", "al/ar/bg x,y,z", "9 x uint32 LE", "milli-g; 0xFFFFFFFF means missing"),
    ], [0.7, 1.6, 1.75, 2.45], 8.2)
    doc.add_heading("6.2 Peak window record: 302 bytes", level=2)
    add_table(doc, ["Offset", "Field", "Wire type", "Meaning"], [
        ("0", "windowStartMm", "int32 LE", "Window start position"),
        ("4", "windowEndMm", "int32 LE", "Window end position"),
        ("8", "speedKmph", "float32 LE", "Maximum speed in window"),
        ("12", "validMask", "uint8", "Node availability mask"),
        ("13", "alertGenerated", "uint8/bool", "Speed gate and threshold exceeded"),
        ("14..301", "9 PeakAxisRecord groups", "9 x 32 bytes", "AL/AR/BG x,y,z"),
    ], [0.7, 1.6, 1.6, 2.6], 8.2)
    add_table(doc, ["PeakAxis offset", "Field", "Type", "Units"], [
        ("+0", "peakValueMg", "uint32 LE", "milli-g"),
        ("+4", "peakPositionMm", "int32 LE", "mm"),
        ("+8", "peakMasterCount", "uint64 LE", "count"),
        ("+16", "peakLat", "float64 LE", "degrees"),
        ("+24", "peakLon", "float64 LE", "degrees"),
    ], [1.2, 1.7, 1.8, 1.8], 8.4)
    doc.add_heading("6.3 Fault record: 75 bytes", level=2)
    add_table(doc, ["Offset", "Field", "Type", "Rule"], [
        ("0", "timestampMs", "uint64 LE", "Clock source is inconsistent; use primarily for ordering"),
        ("8", "faultCode", "uint8", "0x10 timeout, 0x40 upload failed, 0x60 segment invalid, etc."),
        ("9", "nodeId", "uint8", "1 AL, 2 AR, 3 BG, 4 encoder, 0 gateway"),
        ("10", "severity", "uint8", "0 unset, 1 warning, 2 error"),
        ("11", "description", "char[64]", "Null-terminated ASCII; max 63 characters"),
    ], [0.7, 1.4, 1.4, 3.0], 8.2)
    add_callout(doc, "Parser rules", "All fixed records are packed with no padding and use little-endian encoding. Reject unsupported schema versions. Mark a file incomplete if size modulo record size is non-zero. Preserve 0xFFFFFFFF sentinel values or map them to NULL with an explicit quality flag.", AMBER)

    doc.add_heading("7. Cloud processing responsibilities", level=1)
    add_table(doc, ["Stage", "Required processing", "Stored evidence"], [
        ("Receipt", "Stream/buffer ZIP, compute SHA-256, capture size and receive time", "archives"),
        ("Identity", "Compare filename gateway/train/session with metadata", "validation status/error"),
        ("Idempotency", "Check unique (gatewayId, sessionName)", "unique constraint"),
        ("Completeness", "Check 8 required paths and fixed record remainders", "missing_files, integrity_ok"),
        ("Preservation", "Retain original ZIP and extracted files unchanged", "object URI + extracted_files"),
        ("Parsing", "Convert RMS, peaks, faults to structured records", "rms_records, peak_records, fault_records"),
        ("Summary", "Create session/axle projection for dashboard", "gateway_sessions, axle_records"),
        ("Alerting", "Apply route threshold when speed >= 80 km/h; attach GPS/KM", "alerts"),
        ("Notification", "Queue or POST alert notification", "notification_deliveries"),
        ("TMS", "Generate the two required handoff datasets; package for MDB-preferred TMS handoff", "tms_deliveries"),
    ], [1.15, 3.55, 1.8], 8.2)

    doc.add_heading("8. Cloud storage structure", level=1)
    doc.add_paragraph("Use two storage classes. PostgreSQL stores searchable records and relationships. Durable object storage stores immutable ZIP and binary artifacts. A database path alone is not enough for permanent raw archive retention.")
    add_table(doc, ["Storage class", "Content", "Production recommendation"], [
        ("PostgreSQL", "Gateway/session metadata, parsed RMS/peak/fault data, thresholds, alerts, audits", "Managed PostgreSQL with backups, TLS, restricted network, migrations"),
        ("Object storage", "Original ZIP and extracted binaries", "S3/GCS/Azure Blob or railway object store with versioning and retention lock"),
        ("Temporary work area", "Extraction/parsing scratch files", "Encrypted ephemeral disk; delete after durable write and parse"),
        ("TMS output", "Generated handoff package: spatial dataset, processed-peak dataset, README, optional MDB", "Object store or direct CRIS transfer with checksum audit"),
    ], [1.35, 2.55, 2.6], 8.5)
    doc.add_heading("8.1 PostgreSQL logical tables", level=2)
    add_table(doc, ["Table", "Key / relationship", "Purpose"], [
        ("gateways", "gateway_id PK", "Gateway online/offline and last contact"),
        ("trains", "train_id PK", "Train roster"),
        ("archives", "UNIQUE gateway_id + session_name", "Authoritative ZIP receipt metadata/checksum"),
        ("gateway_sessions", "id PK; gateway FK", "Queryable session summary and GPS/speed"),
        ("extracted_files", "archive_id/session_id FK", "Path, size, integrity of every extracted file"),
        ("rms_records", "archive_id/session_id FK", "66-byte records expanded to columns"),
        ("peak_records", "archive_id/session_id FK", "302-byte windows; axes stored as JSON"),
        ("fault_records", "archive_id/session_id FK", "Parsed gateway faults"),
        ("axle_records", "session_id FK", "Dashboard RMS/peak/vertical/lateral summary"),
        ("alerts", "session_id FK", "Threshold violations with route and nearest KM"),
        ("threshold_settings", "route UNIQUE", "Editable route thresholds"),
        ("calibration", "gateway_id + axle_id history", "Wheel diameter/wear/correction history"),
        ("route_track_points", "route index", "GPS-to-nearest-KM reference"),
        ("route_sections", "route index", "Railway/division/section/KM reporting"),
        ("notification_deliveries", "alert_id FK", "SMS/webhook delivery audit"),
        ("tms_deliveries", "id PK", "CRIS package transfer audit"),
    ], [1.65, 2.2, 2.65], 8.1)

    doc.add_page_break()
    doc.add_heading("8.2 Core SQL column types", level=2)
    add_table(doc, ["Entity", "Important columns and data types"], [
        ("archives", "id INTEGER; gateway_id VARCHAR(64); train_id VARCHAR(64); session_name VARCHAR(64); archive_size_bytes INTEGER; received TIMESTAMP; storage_uri VARCHAR(512); checksum VARCHAR(64); validation_status VARCHAR(32); metadata JSON"),
        ("gateway_sessions", "id INTEGER; session_id VARCHAR(64); gateway/train/route VARCHAR; timestamp TIMESTAMP; lat/lon/speed DOUBLE PRECISION; raw_payload JSON"),
        ("rms_records", "master_count/position INTEGER; lat/lon FLOAT; gps_valid BOOLEAN; valid_mask INTEGER; nine acceleration INTEGER columns"),
        ("peak_records", "window positions INTEGER; speed FLOAT; valid_mask INTEGER; alert_generated BOOLEAN; axes JSON"),
        ("alerts", "metric/severity/message VARCHAR; value/threshold/speed/nearest_km FLOAT; created_at TIMESTAMP"),
    ], [1.45, 5.05], 8)

    doc.add_heading("9. What the cloud sends", level=1)
    add_table(doc, ["Destination", "Output", "Format / transport", "Status"], [
        ("Gateway", "Threshold, wear, sampling rate", "GET /api/v1/config JSON over HTTPS", "Cloud endpoint exists; gateway polling not in frozen implementation"),
        ("Officials", "Safety alert with value and GPS", "Webhook JSON to SMS/notification provider", "Outbox works; provider URL/token required"),
        ("Dashboard", "Cards, trends, GPS, alerts", "REST JSON over HTTPS", "Verified HTTP 200"),
        ("CRIS/TMS", "Spatial acceleration + processed peaks", "MDB-preferred handoff package; open ASCII files included where MDB population is not available", "Generator works; CRIS endpoint required"),
    ], [1.15, 2.1, 2.15, 1.1], 8)
    doc.add_heading("9.1 Cloud-to-gateway configuration JSON", level=2)
    add_code(doc, """GET /api/v1/config?gatewayId=GW_BOGIE_001&route=Bangalore-Chennai

{
  "threshold": 50.0,
  "wearPercent": 4.0,
  "samplingRate": 2500
}""")
    add_callout(doc, "Interface decision required", "The current config response has one threshold value. A production schema should version the response and provide separate vertical/lateral and per-axis thresholds, effective time, config revision, and acknowledgement state.", AMBER)
    doc.add_heading("9.2 Recommended realtime alert JSON", level=2)
    add_code(doc, """{
  "schemaVersion": "1.0",
  "gatewayId": "GW_BOGIE_001",
  "trainId": "TRAIN_07",
  "sessionName": "SESSION_20260609_083015",
  "windowStartMm": 250000,
  "windowEndMm": 300000,
  "speedKmph": 92.4,
  "triggeredAxes": [{
    "axisName": "al_z",
    "peakValueMg": 14500,
    "thresholdMg": 12000,
    "peakPositionMm": 276125,
    "peakLat": 13.0827,
    "peakLon": 80.2707
  }]
}""")
    doc.add_heading("9.3 CRIS/TMS handoff package", level=2)
    doc.add_paragraph("The specification does not say the data should be stored in MMD. MMD is a physical clearance envelope. The data-file preference mentioned for TMS is MDB. The same clause also allows storage in a database or ASCII file compatible with TMS. Therefore, the production target should be MDB when CRIS requires it, while the reference cloud keeps PostgreSQL as the live database and includes open ASCII/CSV datasets so the two required data types are documented and importable.")
    add_table(doc, ["File", "Content"], [
        ("spatial_acceleration_export.csv", "RMS spatial records with session, position, GPS, validity and nine axes"),
        ("processed_peak_export.csv", "One row per peak axis/window with position, GPS, speed and alert flag"),
        ("README_MDB_EXPORT.txt", "Schema/import guidance"),
        ("uabams_tms_target.mdb", "Preferred TMS container. On Linux/Render this can be created as a valid empty container; populated MDB requires Windows Jet/ACE import"),
    ], [2.25, 4.25], 8.5)
    add_callout(doc, "MMD clarification", "MMD is not a TMS data file and not a CSV export. Maximum Moving Dimension compliance is a mechanical/hardware installation requirement for accelerometers, system hardware, and mounting brackets within the IR Schedule of Dimension envelope. The cloud can store a compliance document or audit reference, but it cannot prove MMD by itself.", AMBER)

    doc.add_heading("9.4 MMD/SOD compliance evidence for implementers", level=2)
    add_table(doc, ["Evidence item", "Owner", "Cloud role"], [
        ("LHB/FIAT or Vande Bharat coach/bogie drawing reference", "Hardware/vendor team", "Store document reference only if required"),
        ("Sensor and bracket 3D envelope/drawing", "Hardware/vendor team", "Store attachment/checksum only if required"),
        ("Installation location photos and measurements", "Installation team", "Optional audit attachment"),
        ("MMD/SOD clearance calculation and sign-off", "Mechanical approver / railway authority", "Optional compliance metadata"),
        ("Shock/vibration/environment certification", "Hardware/vendor team", "Optional document register"),
    ], [2.3, 1.65, 2.55], 8.2)

    doc.add_heading("10. API catalogue", level=1)
    add_table(doc, ["Method", "Path", "Direction", "Purpose"], [
        ("PUT", "/api/v1/archive", "Gateway -> Cloud", "Production ZIP ingestion"),
        ("POST", "/api/v1/archive", "Tester -> Cloud", "Demo JSON or compatible upload"),
        ("GET", "/api/v1/config", "Cloud -> Gateway", "Threshold/wear/sampling config"),
        ("GET/POST", "/api/v1/threshold", "UI <-> Cloud", "Read/update route limits"),
        ("GET/POST", "/api/v1/calibration", "UI <-> Cloud", "Calibration history"),
        ("GET", "/api/v1/alerts", "Cloud -> UI", "Alert query/filter"),
        ("GET", "/api/v1/dashboard", "Cloud -> UI", "Dashboard aggregate JSON"),
        ("GET", "/api/v1/export/tms", "Cloud -> User", "Download TMS ZIP"),
        ("POST", "/api/v1/export/tms/deliver", "Cloud -> CRIS", "Deliver/audit TMS package"),
        ("GET", "/health and /health/db", "Operations", "Service and database readiness"),
    ], [0.7, 2.25, 1.35, 2.2], 8.1)

    doc.add_heading("11. How to view cloud data", level=1)
    doc.add_heading("11.1 API and dashboard", level=2)
    add_bullet(doc, "Dashboard: https://uabams1-1.onrender.com")
    add_bullet(doc, "Interactive API documentation: https://uabams1.onrender.com/docs")
    add_bullet(doc, "Database health counts: https://uabams1.onrender.com/health/db")
    add_bullet(doc, "Structured endpoints: /api/v1/dashboard, /api/v1/alerts, /api/v1/gateways")
    doc.add_heading("11.2 PostgreSQL viewer", level=2)
    doc.add_paragraph("Use DBeaver, pgAdmin, or psql. DB Browser for SQLite cannot open PostgreSQL. In Render, open uabams-db -> Info -> External Database URL. Create a PostgreSQL connection using the displayed host, port 5432, database, username, password, and SSL=require. Never place the database URL in source code or screenshots.")
    add_code(doc, """SELECT gateway_id, status, last_upload_at FROM gateways ORDER BY gateway_id;

SELECT session_id, gateway_id, train_id, route, timestamp, speed_kmph
FROM gateway_sessions ORDER BY timestamp DESC LIMIT 50;

SELECT gateway_id, train_id, route, metric, value, threshold_value,
       severity, created_at
FROM alerts ORDER BY created_at DESC LIMIT 100;

SELECT a.archive_name, a.validation_status, a.checksum,
       COUNT(r.id) AS rms_rows, COUNT(DISTINCT p.id) AS peak_rows
FROM archives a
LEFT JOIN rms_records r ON r.archive_id = a.id
LEFT JOIN peak_records p ON p.archive_id = a.id
GROUP BY a.id ORDER BY a.upload_received_utc DESC;""")

    doc.add_heading("12. Security and operations requirements", level=1)
    add_table(doc, ["Control", "Production requirement"], [
        ("Transport", "TLS 1.2+; disable plain HTTP; approved certificate chain"),
        ("Gateway authentication", "Per-device token or mutual TLS; rotate/revoke credentials"),
        ("Authorization", "Gateway may upload only for its own gatewayId; operator roles for settings/export"),
        ("Secrets", "Environment secret store; never commit database/API credentials"),
        ("Database", "Private network, TLS, least privilege, backups, point-in-time recovery"),
        ("Object storage", "Encryption, versioning, retention policy, checksum verification"),
        ("Observability", "Structured logs, upload latency/error metrics, alerts for repeated failures"),
        ("Schema evolution", "Versioned JSON/binary schemas and database migrations (Alembic)"),
        ("Time", "UTC everywhere; retain original timestamp and clock-quality flag"),
        ("Privacy", "Restrict GPS/train operational data; audit access and exports"),
    ], [1.55, 4.95], 8.5)

    doc.add_heading("13. Requirement-to-implementation matrix", level=1)
    add_table(doc, ["Manager / ICD requirement", "Reference status", "Railman implementation action"], [
        ("Document data structure and types", "Covered in Sections 5, 6, 8", "Implement identical versioned parsers and DB mapping"),
        ("Define what cloud receives/sends", "Covered in Sections 5 and 9", "Freeze endpoint URLs, auth, and ownership"),
        ("JSON and binary formats", "Metadata and demo JSON plus packed records documented", "Use metadata v1.0; test golden binary fixtures"),
        ("Gateway connects to cloud", "HTTP PUT contract implemented", "Use HTTPS, Content-Length, device identity, retries"),
        ("Database/ASCII storage", "PostgreSQL database implemented; open ASCII export included for compatibility", "Configure backups and migration ownership"),
        ("Spatial acceleration", "rms_records parser implemented", "Prove with real rms_25cm.bin"),
        ("Processed peaks", "peak_records parser implemented", "Prove all nine axes and sentinel cases"),
        ("GPS safety alerts", "Archive rule engine and UI implemented", "Complete realtime gateway POST and SMS provider"),
        ("Route-wise thresholds", "Implemented in cloud UI/DB", "Agree config schema and gateway acknowledgement"),
        ("Permanent raw retention", "Code writes files, but current Render disk is not durable", "Add object storage before production"),
        ("CRIS/TMS transfer", "Two required datasets packaged for MDB-preferred handoff; HTTP delivery code implemented", "Obtain final CRIS protocol, schema, URL and credentials"),
        ("MDB preference", "MDB is the preferred TMS container; populated MDB requires Windows Jet/ACE or final CRIS import process", "Use MDB as final target when CRIS confirms schema/tooling"),
        ("MMD/SOD compliance", "Not a cloud data export; documented as hardware evidence", "Hardware team must prove envelope compliance from drawings/measurements"),
    ], [2.35, 1.8, 2.35], 7.8)

    doc.add_heading("14. Acceptance test checklist", level=1)
    for item in (
        "Upload a known closed-session ZIP and verify HTTP 200 only after object-store write.",
        "Re-upload the same ZIP and verify no duplicate rows and HTTP 200.",
        "Change filename gatewayId and verify HTTP 422 plus quarantine evidence.",
        "Remove one required file and verify incomplete status while preserving the archive.",
        "Truncate each fixed-record file and verify integrity failure is recorded.",
        "Verify 0xFFFFFFFF sentinel and gpsValid=false handling.",
        "Verify nine peak axes, speed gate, route thresholds, GPS and nearest KM.",
        "Verify SMS/webhook success and failure audit rows.",
        "Verify cloud config authentication, revision, and gateway acknowledgement.",
        "Download and validate both TMS handoff datasets and SHA-256 audit.",
        "Restore PostgreSQL and object data from backup in a recovery test.",
    ):
        add_bullet(doc, "[ ] " + item)

    doc.add_heading("15. Final handoff statement", level=1)
    doc.add_paragraph("Railman can use this document as the cloud implementation contract, subject to four interface decisions: production authentication, durable object-storage platform, final realtime alert endpoint/schema, and final CRIS/TMS transport/schema. The supplied reference application demonstrates the core parsing, PostgreSQL, dashboard, alert, threshold, calibration, and export concepts; it should be treated as an executable reference, not as a completed railway production acceptance.")

    doc.add_page_break()
    doc.add_heading("Appendix A. Data ownership summary", level=1)
    add_table(doc, ["Data", "System of record", "Retention", "Primary consumer"], [
        ("Original session ZIP", "Object storage", "Permanent / policy-defined immutable", "Audit and replay"),
        ("Extracted binaries", "Object storage", "Permanent", "Engineering/forensics"),
        ("Parsed RMS/peak/fault", "PostgreSQL", "Operational + archive policy", "Analytics/dashboard/TMS"),
        ("Threshold/calibration", "PostgreSQL", "Version history", "Cloud and gateway configuration"),
        ("Alerts/notifications", "PostgreSQL", "Safety/audit policy", "Officials and dashboard"),
        ("TMS handoff packages", "Object storage + audit table", "Handoff policy", "CRIS TMS"),
        ("MMD/SOD evidence", "Document register if required", "Project/contract policy", "Mechanical compliance reviewers"),
    ], [1.55, 1.55, 1.8, 1.6], 8.2)

    doc.add_heading("Appendix B. Source basis", level=1)
    doc.add_paragraph("This handoff was prepared from the UABAMS Cloud Interface Document v1.2 (Frozen, 09 June 2026), the supplied RDSO technical specification excerpt, and the current repository implementation in backend/app. Live endpoint checks were performed on 18 June 2026. Where the interface document and reference implementation differ, the difference is explicitly identified as a gap or decision.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
