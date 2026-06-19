from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "UABAMS_Cloud_Quick_Handoff_3_Pages.docx"
FLOW = ROOT / "_cloud_data_flow.png"
CLOUD_DATA_SCREEN = Path(r"C:\Users\Pilabs\AppData\Local\Temp\codex-clipboard-c00d8937-d6e9-4883-8246-e5f7b33978ce.png")

NAVY = "123A63"
BLUE = "2E74B5"
LIGHT = "F4F6F9"
PALE = "E8EEF5"
AMBER = "FFF4D6"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in (("top", 70), ("start", 90), ("bottom", 70), ("end", 90)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def table(doc, headers, rows, widths, font_size=7.8):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], PALE)
        p = t.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = rgb(NAVY)
    for row in rows:
        tr = t.add_row()
        for i, value in enumerate(row):
            cell = tr.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            r.font.size = Pt(font_size)
    grid = t._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(w * 1440)))
        grid.append(col)
    for row in t.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i])
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(int(widths[i] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            margins(cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def callout(doc, text, fill=AMBER):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    shade(t.cell(0, 0), fill)
    p = t.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.font.color.rgb = rgb(NAVY)
    margins(t.cell(0, 0))


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.space_after = Pt(4)
    for line in text.strip().splitlines():
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r.font.size = Pt(7.2)
        r.font.color.rgb = rgb("1F2933")


def configure(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.45)
    sec.bottom_margin = Inches(0.45)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(8.7)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0
    for style_name, size, color in (
        ("Heading 1", 13, BLUE),
        ("Heading 2", 10.5, NAVY),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(5)
        style.paragraph_format.space_after = Pt(3)


def build():
    doc = Document()
    configure(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("UABAMS Cloud Handoff - 3 Page Review Note")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = rgb(NAVY)
    doc.add_paragraph("Purpose: Give Railman/cloud implementers the minimum cloud data contract, storage model, and workflow needed to build their own cloud review implementation.")

    callout(doc, "Key correction: MDB is the preferred TMS data container. MMD is only the physical Maximum Moving Dimension mounting envelope, not a cloud data format.")

    doc.add_heading("1. Cloud Workflow", level=1)
    if FLOW.exists():
        doc.add_picture(str(FLOW), width=Inches(6.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table(doc, ["Step", "Cloud role", "Output"], [
        ("1", "Receive closed-session gateway archive over HTTPS", "Archive accepted/rejected"),
        ("2", "Validate metadata, gateway ID, train ID, session, checksum, required files", "Validation audit"),
        ("3", "Parse spatial acceleration, processed peak, fault and GPS data", "Database records"),
        ("4", "Apply route thresholds and speed gate; attach GPS/KM context", "Alert rows"),
        ("5", "Serve dashboard/API and prepare MDB-preferred TMS handoff", "Review UI + TMS ZIP"),
    ], [0.45, 4.2, 2.1])

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("2. What Cloud Receives, Stores, Processes, Sends", level=1)
    table(doc, ["Area", "Required structure / data type"], [
        ("Gateway connection", "HTTPS PUT/POST to /api/v1/archive. Body is a closed session archive or demo JSON during review. Production uses bearer token/API key; mTLS can be added by deployer."),
        ("Received metadata JSON", "schemaVersion string, gatewayId string, trainId string, sessionName string, sessionStatus string, createdUtc ISO datetime."),
        ("Received measurement data", "Spatial acceleration records, processed peak records, GPS latitude/longitude, speed, axle ID, metric, timestamp, fault records."),
        ("Cloud storage", "PostgreSQL tables: gateways, gateway_sessions, axle_records, alerts, threshold_settings, calibration, archives, rms_records, peak_records, fault_records, notification_deliveries, tms_deliveries."),
        ("Processing", "Authenticate request, validate upload, parse records, calculate dashboard summaries, compare peak values with route-wise vertical/lateral thresholds, generate GPS alerts."),
        ("Cloud sends to gateway", "Configuration JSON: route threshold, wheel wear/correction factor, sampling rate. Gateway polling/ack should be finalized in production."),
        ("Cloud sends to SMS server", "When an alert is generated, cloud posts SMS JSON to SMS_SERVER_URL and stores success/failure in notification_deliveries."),
        ("Cloud sends to TMS/CRIS", "MDB-preferred handoff package containing spatial acceleration data and processed peak data. ASCII text files are included only as open documented fallback/import data."),
    ], [1.65, 5.1], 7.5)

    doc.add_heading("3. JSON/Data Format Examples", level=1)
    code(doc, r'''
Gateway demo JSON:
{
  "gatewayId": "GW001", "trainId": "TRAIN21",
  "route": "Bangalore-Chennai",
  "gps": {"lat": 13.158, "lon": 77.732},
  "speedKmph": 123,
  "axleData": [{"axleId": "AX01", "verticalG": 77.2, "lateralG": 35.1,
                "rms": 27.0, "peak": 77.2}]
}

Cloud-to-gateway config JSON:
{"gatewayId":"GW001","route":"Bangalore-Chennai","verticalThreshold":50,
 "lateralThreshold":80,"samplingRate":2500,"correctionFactor":1.02}

SMS server JSON:
{"to":["+91xxxxxxxxxx"],"message":"UABAMS Info: vertical 77.2g...",
 "alert":{"gatewayId":"GW004","trainId":"TRAIN21","gps":{"lat":13.15,"lon":77.73}}}
''')

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("4. Screenshots To Add For Review", level=1)
    if CLOUD_DATA_SCREEN.exists():
        doc.add_picture(str(CLOUD_DATA_SCREEN), width=Inches(6.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table(doc, ["Screenshot", "Where to place", "Why it helps Railman"], [
        ("Cloud Data page", "Page 3, under this section", "Shows all cloud API/database data in one place."),
        ("Dashboard overview", "Optional appendix/email body", "Shows graphs, GPS positions, sessions and TMS export button."),
        ("Gateway Upload", "Optional appendix/email body", "Shows how gateway data enters the cloud."),
        ("Threshold Settings", "Optional appendix/email body", "Shows route-wise editable limits from the specification."),
        ("Alerts page", "Optional appendix/email body", "Shows GPS alert map, trend and alert table."),
    ], [1.7, 1.8, 3.25], 7.3)

    doc.add_heading("5. Can Railman Implement From This?", level=1)
    doc.add_paragraph("Yes, for cloud review/prototype implementation. This document gives the cloud workflow, received/sent data, storage structure, JSON examples, processing responsibilities and gateway connection type.")
    doc.add_paragraph("Before production, Railman still must finalize: exact MDB schema/import method with CRIS, durable object storage for archives, SMS provider credentials/recipients, and hardware MMD/SOD compliance evidence.")
    table(doc, ["Production env variable", "Purpose"], [
        ("DATABASE_URL", "PostgreSQL cloud database"),
        ("API_AUTH_TOKEN", "Operator/frontend API authentication"),
        ("GATEWAY_API_TOKEN", "Gateway archive/config authentication"),
        ("VITE_API_TOKEN", "Frontend token matching API_AUTH_TOKEN"),
        ("SMS_SERVER_URL", "SMS server endpoint for alert messages"),
        ("SMS_SERVER_BEARER_TOKEN", "SMS server authentication token"),
        ("SMS_RECIPIENTS", "Comma-separated officials/recipient numbers"),
    ], [2.2, 4.55], 7.2)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
